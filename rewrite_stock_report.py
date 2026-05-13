from __future__ import annotations

import math
import shutil
import zipfile
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.shared import Inches
from PIL import Image, ImageDraw, ImageFont


SOURCE_DOC = Path(r"C:\Users\sarth\Downloads\microproject_UIUX (1).docx")
OUTPUT_DOC = Path(r"C:\Users\sarth\Downloads\Stock_Buy_Sell_App_Project_Report.docx")
ASSET_DIR = Path(r"C:\Users\sarth\OneDrive\Desktop\ai\stock_report_assets")


FONT_REGULAR = Path(r"C:\Windows\Fonts\arial.ttf")
FONT_BOLD = Path(r"C:\Windows\Fonts\arialbd.ttf")


TEXT_REPLACEMENTS = {
    6: '"Stock Buy and Sell App"',
    11: "Narode sir",
    13: "SUBMITTED BY :",
    14: "1. Sarthak Bhokare (26)",
    15: "2. Piyush Borse (29)",
    16: "3. Chine Roshan (33)",
    17: "4. Aditya Chavan (30)",
    20: "INTRODUCTION",
    22: (
        "The Stock Buy and Sell App is a digital trading platform designed to simplify the process of "
        "investing in shares through a mobile interface. Traditional stock trading can feel complicated "
        "for many users because it involves scattered information, confusing workflows, and delayed "
        "decision-making. This project provides a modern solution through a user-friendly interface "
        "designed using Figma."
    ),
    24: (
        "The main aim of the project is to design an intuitive, efficient, and visually appealing UI/UX "
        "that enhances user experience and ensures seamless navigation in a trading environment. The "
        "system allows users to explore stocks, analyze charts, add funds, place buy or sell orders, and "
        "manage their portfolio digitally. This project focuses on creating UI/UX designs using Figma to "
        "ensure clarity, trust, and ease of use for both beginners and active investors. With the growing "
        "use of smartphones and digital finance platforms, users expect quick and secure services. A "
        "well-designed stock buy and sell app helps users compare prices, study market trends, and take "
        "investment decisions without confusion or delay."
    ),
    25: "OBJECTIVES OF THE PROJECT",
    28: "To design a simple, user-friendly, and intuitive trading interface",
    29: "To minimize user effort and reduce the time required to place an order",
    30: "To create a smooth and logical user flow",
    31: "To enhance user satisfaction and engagement",
    32: "To implement modern UI/UX principles such as consistency and feedback",
    33: "To design a system that is accessible for both beginners and active traders",
    34: "To ensure clarity in charts, prices, and order actions",
    35: "To provide a visually attractive and professional financial dashboard",
    36: "To reduce confusion during buying and selling operations",
    37: "To support fast and secure trading interactions",
    41: "SCOPE OF THE PROJECT",
    43: "The scope of the project is limited to designing the UI/UX of the stock buy and sell app using Figma.",
    44: "It includes:",
    46: "Designing the complete user journey from login to order confirmation",
    48: "Creating interactive prototypes for real-time trading flow simulation",
    50: "Designing mobile-friendly and scalable layouts",
    52: "Applying design standards and usability principles",
    54: (
        "However, the project does not include backend development, real-time market integration, payment "
        "gateway processing, or actual trade execution functionality. The focus is entirely on improving "
        "the user experience through effective design."
    ),
    55: "TOOLS USED",
    57: "Figma - For UI design, prototyping, and collaboration",
    59: "Wireframing Tools - For creating initial layouts",
    61: "Design Principles - Color theory, typography, spacing, alignment",
    63: "Icon Libraries - For better visual representation",
    65: "Prototyping Features - To simulate real user interaction",
    66: "Why Figma?",
    67: "Cloud-based and easy to access",
    68: "Supports teamwork and sharing",
    69: "Provides real-time design updates",
    70: "SYSTEM FEATURES",
    72: "1. Login & Signup Screen",
    73: "Clean and secure onboarding interface",
    74: "Easy input fields and clear buttons",
    75: "Options for password recovery and OTP verification",
    76: "Evaluation:",
    77: "Improves user onboarding",
    78: "Reduces confusion for new investors",
    79: "Ensures secure access",
    83: "2. Home Dashboard",
    84: "Market overview with indices, top gainers, and watchlist",
    85: "Quick access to search, portfolio, and trading actions",
    86: "Evaluation:",
    87: "Saves time and effort",
    88: "Provides direct access to core features",
    89: "Improves usability",
    92: "3. Stock Listing Screen",
    93: "Displays available stocks with filters and sector categories",
    94: "Sorting options based on price, change, and volume",
    95: "Evaluation:",
    96: "Helps users compare choices easily",
    97: "Improves decision-making",
    98: "Enhances user controls",
    103: "4. Stock Details Screen",
    104: "Interactive chart with key market metrics",
    105: "Clear indication of price movement and stock actions",
    106: "Evaluation:",
    107: "Provides better visualization",
    108: "Reduces trading mistakes",
    109: "Improves user confidence",
    113: "5. Buy Order Screen",
    114: "Allows user to select quantity, order type, and price",
    115: "Provides estimated cost and brokerage summary",
    116: "Simple and clean order placement interface",
    117: "Evaluation:",
    118: "Improves order accuracy",
    119: "Maintains smooth user flow",
    120: "Adds flexibility for users",
    125: "6. Sell Order Screen",
    126: "Simple form design with holding visibility",
    127: "Easy data entry for quantity and execution choice",
    128: "Ensures security and efficiency",
    130: "Evaluation:",
    131: "Ensures correct order input",
    132: "Reduces complexity",
    133: "Helps in fast transaction processing",
    136: "7. Add Funds / Payment Screen",
    137: "Multiple payment options such as UPI, card, and net banking",
    138: "Secure and simple interface",
    139: "Evaluation:",
    140: "Builds trust among users",
    141: "Ensures smooth transactions",
    145: "8. Order Confirmation Screen",
    146: "Displays order details clearly",
    147: "Option to download or save transaction summary",
    149: "Evaluation:",
    150: "Provides clarity and confirmation",
    151: "Enhances user satisfaction",
    155: "9. Portfolio & Order History Screen",
    156: "Displays holdings, returns, and order history",
    157: "Options to review, cancel, or modify pending orders",
    159: "Evaluation:",
    160: "Improves convenience",
    161: "Provides centralized management",
    165: "ADVANTAGES OF THE SYSTEM",
    167: "Saves time and effort",
    168: "Easy and convenient investing experience",
    169: "Reduces human errors during order placement",
    170: "Accessible anytime and anywhere",
    171: "Improves overall user experience",
    172: "User-friendly for beginners and active traders",
    173: "Supports quick financial decision-making",
    176: "LIMITATION OF THE SYSTEM",
    178: "No backend implementation",
    179: "No live market data integration",
    180: "Limited to design only",
    181: "Requires internet for usage",
    182: "Cannot process actual transactions",
    184: "Future enhancements",
    187: "Integration with real-time market database",
    188: "AI-based stock recommendations",
    189: "Voice-enabled stock search",
    190: "Multi-language support",
    191: "Dark mode UI",
    192: "Mobile app development",
    193: "Personalized alerts and notifications",
    194: "Improved tablet responsiveness",
    195: "Secure biometric login",
    196: "News and sentiment analysis",
    203: "CONCLUSION",
    205: (
        "The Stock Buy and Sell App UI/UX project successfully demonstrates the importance of "
        "user-centered design in financial technology. By focusing on usability, accessibility, visual "
        "hierarchy, and trust, the system provides a seamless trading experience. The use of Figma enables "
        "effective prototyping and testing, ensuring a well-structured design before development."
    ),
    207: (
        "This project highlights how a well-designed interface can improve investor confidence, reduce "
        "friction in transactions, and support better decision-making. With further enhancements and "
        "backend integration, this system can be transformed into a fully functional real-world trading "
        "application."
    ),
}


def load_font(size: int, bold: bool = False) -> ImageFont.FreeTypeFont | ImageFont.ImageFont:
    path = FONT_BOLD if bold else FONT_REGULAR
    try:
        return ImageFont.truetype(str(path), size=size)
    except OSError:
        return ImageFont.load_default()


def vertical_gradient(width: int, height: int, top: tuple[int, int, int], bottom: tuple[int, int, int]) -> Image.Image:
    image = Image.new("RGB", (width, height), top)
    draw = ImageDraw.Draw(image)
    for y in range(height):
        mix = y / max(height - 1, 1)
        color = tuple(int(top[i] * (1 - mix) + bottom[i] * mix) for i in range(3))
        draw.line((0, y, width, y), fill=color)
    return image


def draw_soft_circle(draw: ImageDraw.ImageDraw, center: tuple[int, int], radius: int, fill: tuple[int, int, int]) -> None:
    x, y = center
    draw.ellipse((x - radius, y - radius, x + radius, y + radius), fill=fill)


def text(draw: ImageDraw.ImageDraw, xy: tuple[int, int], value: str, size: int, fill: str = "white", bold: bool = False) -> None:
    draw.text(xy, value, font=load_font(size, bold=bold), fill=fill)


def draw_phone_mockup(
    base: Image.Image,
    title: str,
    subtitle: str,
    variant: str,
    accent: tuple[int, int, int],
) -> None:
    draw = ImageDraw.Draw(base)
    width, height = base.size

    phone_w, phone_h = 430, 820
    phone_x = width // 2 - phone_w // 2
    phone_y = 110
    radius = 46

    draw.rounded_rectangle((phone_x, phone_y, phone_x + phone_w, phone_y + phone_h), radius=radius, fill=(21, 29, 45))
    draw.rounded_rectangle((phone_x + 18, phone_y + 18, phone_x + phone_w - 18, phone_y + phone_h - 18), radius=36, fill=(246, 249, 255))
    draw.rounded_rectangle((phone_x + 145, phone_y + 24, phone_x + 285, phone_y + 44), radius=10, fill=(15, 21, 34))

    inner_x = phone_x + 34
    inner_y = phone_y + 58
    inner_w = phone_w - 68
    bottom = phone_y + phone_h - 38

    text(draw, (90, 78), "UI/UX MICROPROJECT", 28, fill="#E6EFFB", bold=True)
    text(draw, (90, 118), title, 54, fill="white", bold=True)
    text(draw, (90, 174), subtitle, 24, fill="#D2DBE8")

    # Decorative background shapes.
    draw_soft_circle(draw, (220, 840), 180, (32, 56, 89))
    draw_soft_circle(draw, (1330, 160), 120, (35, 79, 116))
    draw_soft_circle(draw, (1240, 830), 210, (29, 51, 79))
    draw.rounded_rectangle((104, 228, 350, 300), radius=26, fill=(255, 255, 255, 20))

    # Phone status and header.
    text(draw, (inner_x, inner_y), "9:41", 18, fill="#6B7280", bold=True)
    text(draw, (inner_x + inner_w - 74, inner_y), "5G", 18, fill="#6B7280", bold=True)
    text(draw, (inner_x, inner_y + 42), "TradeSphere", 24, fill="#111827", bold=True)
    text(draw, (inner_x + inner_w - 100, inner_y + 42), "Live", 20, fill="#16A34A", bold=True)

    card_top = inner_y + 92
    draw.rounded_rectangle((inner_x, card_top, inner_x + inner_w, card_top + 116), radius=26, fill=(17, 24, 39))
    text(draw, (inner_x + 22, card_top + 18), "Portfolio Value", 18, fill="#CBD5E1")
    text(draw, (inner_x + 22, card_top + 48), "Rs. 4,85,240", 34, fill="white", bold=True)
    text(draw, (inner_x + inner_w - 128, card_top + 50), "+3.84%", 28, fill="#4ADE80", bold=True)

    if variant == "login":
        draw.rounded_rectangle((inner_x, card_top + 148, inner_x + inner_w, card_top + 565), radius=28, fill=(255, 255, 255))
        text(draw, (inner_x + 24, card_top + 176), "Welcome back", 32, fill="#111827", bold=True)
        text(draw, (inner_x + 24, card_top + 220), "Invest smarter every day", 20, fill="#6B7280")
        for idx, label in enumerate(["Email or mobile", "Password"]):
            y = card_top + 278 + idx * 110
            text(draw, (inner_x + 24, y), label, 18, fill="#6B7280")
            draw.rounded_rectangle((inner_x + 24, y + 30, inner_x + inner_w - 24, y + 86), radius=18, outline="#D1D5DB", width=2)
        draw.rounded_rectangle((inner_x + 24, card_top + 520, inner_x + inner_w - 24, card_top + 584), radius=22, fill=accent)
        text(draw, (inner_x + 122, card_top + 536), "Login", 24, fill="white", bold=True)
    elif variant == "dashboard":
        draw.rounded_rectangle((inner_x, card_top + 148, inner_x + inner_w, card_top + 278), radius=24, fill=(236, 253, 245))
        text(draw, (inner_x + 20, card_top + 168), "NIFTY 50", 18, fill="#166534", bold=True)
        text(draw, (inner_x + 20, card_top + 198), "22,421.80", 32, fill="#111827", bold=True)
        text(draw, (inner_x + 20, card_top + 236), "+184.25 today", 20, fill="#16A34A", bold=True)
        for idx, symbol in enumerate(["TCS", "INFY", "RELIANCE"]):
            y = card_top + 310 + idx * 92
            draw.rounded_rectangle((inner_x, y, inner_x + inner_w, y + 74), radius=20, fill=(255, 255, 255))
            text(draw, (inner_x + 18, y + 16), symbol, 22, fill="#111827", bold=True)
            text(draw, (inner_x + 18, y + 42), "Watchlist", 16, fill="#6B7280")
            text(draw, (inner_x + inner_w - 136, y + 18), f"Rs. {1480 + idx * 320}", 22, fill="#111827", bold=True)
            text(draw, (inner_x + inner_w - 96, y + 44), "+1.8%", 18, fill="#16A34A", bold=True)
    elif variant == "listing":
        text(draw, (inner_x, card_top + 150), "Market Movers", 28, fill="#111827", bold=True)
        for idx, item in enumerate([("HDFCBANK", "+2.4%"), ("SBIN", "+1.7%"), ("ITC", "-0.8%"), ("WIPRO", "+3.1%")]):
            y = card_top + 198 + idx * 96
            draw.rounded_rectangle((inner_x, y, inner_x + inner_w, y + 78), radius=18, fill=(255, 255, 255))
            text(draw, (inner_x + 18, y + 14), item[0], 20, fill="#111827", bold=True)
            text(draw, (inner_x + 18, y + 42), "NSE", 16, fill="#6B7280")
            color = "#16A34A" if item[1].startswith("+") else "#DC2626"
            text(draw, (inner_x + inner_w - 90, y + 26), item[1], 20, fill=color, bold=True)
    elif variant == "details":
        draw.rounded_rectangle((inner_x, card_top + 148, inner_x + inner_w, card_top + 410), radius=24, fill=(255, 255, 255))
        text(draw, (inner_x + 20, card_top + 170), "RELIANCE", 26, fill="#111827", bold=True)
        text(draw, (inner_x + 20, card_top + 208), "Rs. 2,943.40", 30, fill="#111827", bold=True)
        text(draw, (inner_x + 204, card_top + 210), "+2.15%", 22, fill="#16A34A", bold=True)
        points = []
        for step in range(9):
            x = inner_x + 24 + step * 40
            y = card_top + 342 - int(math.sin(step / 1.4) * 62 + step * 7)
            points.append((x, y))
        draw.line(points, fill=accent, width=6, joint="curve")
        for px, py in points:
            draw.ellipse((px - 6, py - 6, px + 6, py + 6), fill=accent)
        draw.rounded_rectangle((inner_x, card_top + 438, inner_x + 160, card_top + 506), radius=22, fill=accent)
        draw.rounded_rectangle((inner_x + 182, card_top + 438, inner_x + inner_w, card_top + 506), radius=22, fill=(17, 24, 39))
        text(draw, (inner_x + 44, card_top + 458), "Buy", 24, fill="white", bold=True)
        text(draw, (inner_x + 292, card_top + 458), "Sell", 24, fill="white", bold=True)
    elif variant == "buy":
        text(draw, (inner_x, card_top + 150), "Buy TCS", 30, fill="#111827", bold=True)
        for idx, label in enumerate(["Quantity", "Order Type", "Limit Price"]):
            y = card_top + 212 + idx * 104
            text(draw, (inner_x, y), label, 18, fill="#6B7280")
            draw.rounded_rectangle((inner_x, y + 28, inner_x + inner_w, y + 84), radius=18, outline="#CBD5E1", width=2)
        draw.rounded_rectangle((inner_x, card_top + 558, inner_x + inner_w, card_top + 650), radius=24, fill=(239, 246, 255))
        text(draw, (inner_x + 22, card_top + 584), "Estimated Cost", 18, fill="#1D4ED8")
        text(draw, (inner_x + 22, card_top + 614), "Rs. 17,982", 28, fill="#111827", bold=True)
        draw.rounded_rectangle((inner_x, card_top + 680, inner_x + inner_w, card_top + 748), radius=22, fill=accent)
        text(draw, (inner_x + 110, card_top + 698), "Proceed to Buy", 24, fill="white", bold=True)
    elif variant == "sell":
        text(draw, (inner_x, card_top + 150), "Sell Infosys", 30, fill="#111827", bold=True)
        draw.rounded_rectangle((inner_x, card_top + 204, inner_x + inner_w, card_top + 292), radius=22, fill=(254, 242, 242))
        text(draw, (inner_x + 20, card_top + 226), "Available Holdings", 18, fill="#991B1B")
        text(draw, (inner_x + 20, card_top + 256), "32 Shares", 28, fill="#111827", bold=True)
        for idx, label in enumerate(["Quantity to Sell", "Order Type", "Target Price"]):
            y = card_top + 334 + idx * 96
            text(draw, (inner_x, y), label, 18, fill="#6B7280")
            draw.rounded_rectangle((inner_x, y + 28, inner_x + inner_w, y + 82), radius=18, outline="#CBD5E1", width=2)
        draw.rounded_rectangle((inner_x, card_top + 664, inner_x + inner_w, card_top + 732), radius=22, fill=(220, 38, 38))
        text(draw, (inner_x + 116, card_top + 682), "Review Sell Order", 24, fill="white", bold=True)
    elif variant == "funds_upi":
        text(draw, (inner_x, card_top + 150), "Add Funds", 30, fill="#111827", bold=True)
        for idx, label in enumerate(["UPI", "Net Banking", "Debit Card"]):
            y = card_top + 212 + idx * 100
            draw.rounded_rectangle((inner_x, y, inner_x + inner_w, y + 78), radius=20, fill=(255, 255, 255))
            text(draw, (inner_x + 20, y + 22), label, 22, fill="#111827", bold=True)
            text(draw, (inner_x + inner_w - 108, y + 24), "Select", 20, fill="#2563EB", bold=True)
        draw.rounded_rectangle((inner_x, card_top + 546, inner_x + inner_w, card_top + 700), radius=24, fill=(17, 24, 39))
        text(draw, (inner_x + 22, card_top + 570), "Amount", 18, fill="#CBD5E1")
        text(draw, (inner_x + 22, card_top + 606), "Rs. 25,000", 34, fill="white", bold=True)
        draw.rounded_rectangle((inner_x + 22, card_top + 648, inner_x + inner_w - 22, card_top + 690), radius=18, fill=accent)
        text(draw, (inner_x + 124, card_top + 656), "Continue", 22, fill="white", bold=True)
    elif variant == "funds_card":
        text(draw, (inner_x, card_top + 150), "Payment Method", 30, fill="#111827", bold=True)
        draw.rounded_rectangle((inner_x, card_top + 214, inner_x + inner_w, card_top + 410), radius=26, fill=(37, 99, 235))
        text(draw, (inner_x + 24, card_top + 244), "Virtual Trade Card", 24, fill="white", bold=True)
        text(draw, (inner_x + 24, card_top + 304), "**** 7045", 36, fill="white", bold=True)
        text(draw, (inner_x + 24, card_top + 352), "Expires 09/29", 18, fill="#DBEAFE")
        for idx, label in enumerate(["UPI AutoPay", "Saved Cards", "Net Banking"]):
            y = card_top + 448 + idx * 88
            draw.rounded_rectangle((inner_x, y, inner_x + inner_w, y + 68), radius=18, fill=(255, 255, 255))
            text(draw, (inner_x + 18, y + 18), label, 20, fill="#111827", bold=True)
        draw.rounded_rectangle((inner_x, card_top + 720, inner_x + inner_w, card_top + 788), radius=22, fill=accent)
        text(draw, (inner_x + 120, card_top + 738), "Pay Securely", 24, fill="white", bold=True)
    elif variant == "funds_success":
        draw.rounded_rectangle((inner_x, card_top + 170, inner_x + inner_w, card_top + 640), radius=30, fill=(255, 255, 255))
        draw.ellipse((inner_x + 118, card_top + 220, inner_x + 278, card_top + 380), fill=(34, 197, 94))
        draw.line((inner_x + 162, card_top + 304, inner_x + 194, card_top + 336), fill="white", width=10)
        draw.line((inner_x + 194, card_top + 336, inner_x + 246, card_top + 268), fill="white", width=10)
        text(draw, (inner_x + 62, card_top + 422), "Funds Added Successfully", 28, fill="#111827", bold=True)
        text(draw, (inner_x + 126, card_top + 468), "Rs. 25,000 credited", 22, fill="#16A34A", bold=True)
        text(draw, (inner_x + 74, card_top + 522), "Ready for your next trade", 20, fill="#6B7280")
        draw.rounded_rectangle((inner_x + 40, card_top + 566, inner_x + inner_w - 40, card_top + 624), radius=20, fill=accent)
        text(draw, (inner_x + 96, card_top + 582), "Back to Trading", 24, fill="white", bold=True)
    elif variant == "confirmation":
        draw.rounded_rectangle((inner_x, card_top + 172, inner_x + inner_w, card_top + 670), radius=28, fill=(255, 255, 255))
        text(draw, (inner_x + 24, card_top + 206), "Order Confirmed", 30, fill="#111827", bold=True)
        text(draw, (inner_x + 24, card_top + 252), "BUY 12 shares of TCS", 22, fill="#16A34A", bold=True)
        for idx, row in enumerate([("Order ID", "#TRD-24081"), ("Amount", "Rs. 44,184"), ("Order Type", "Market Order"), ("Status", "Completed")]):
            y = card_top + 314 + idx * 76
            text(draw, (inner_x + 24, y), row[0], 18, fill="#6B7280")
            text(draw, (inner_x + inner_w - 170, y), row[1], 20, fill="#111827", bold=True)
        draw.rounded_rectangle((inner_x + 24, card_top + 620, inner_x + inner_w - 24, card_top + 678), radius=20, fill=accent)
        text(draw, (inner_x + 86, card_top + 636), "Download Summary", 22, fill="white", bold=True)
    elif variant == "portfolio":
        text(draw, (inner_x, card_top + 150), "Portfolio", 30, fill="#111827", bold=True)
        draw.rounded_rectangle((inner_x, card_top + 204, inner_x + inner_w, card_top + 332), radius=24, fill=(240, 253, 244))
        text(draw, (inner_x + 20, card_top + 228), "Today's Return", 18, fill="#166534")
        text(draw, (inner_x + 20, card_top + 260), "+Rs. 8,420", 34, fill="#111827", bold=True)
        text(draw, (inner_x + 20, card_top + 302), "+1.96%", 22, fill="#16A34A", bold=True)
        for idx, symbol in enumerate(["TCS", "RELIANCE", "HDFCBANK"]):
            y = card_top + 370 + idx * 96
            draw.rounded_rectangle((inner_x, y, inner_x + inner_w, y + 78), radius=18, fill=(255, 255, 255))
            text(draw, (inner_x + 18, y + 16), symbol, 20, fill="#111827", bold=True)
            text(draw, (inner_x + 18, y + 44), "Qty 12", 16, fill="#6B7280")
            text(draw, (inner_x + inner_w - 128, y + 28), "+4.2%", 20, fill="#16A34A", bold=True)

    text(draw, (90, 908), "Stock trading app concept screens", 22, fill="#C7D2E1")
    draw.rounded_rectangle((0, bottom, width, bottom + 8), radius=0, fill=accent)


def create_screen_image(path: Path, title: str, subtitle: str, variant: str, accent: tuple[int, int, int]) -> None:
    base = vertical_gradient(1600, 1000, (12, 25, 46), (25, 78, 111))
    draw_phone_mockup(base, title, subtitle, variant, accent)
    base.save(path, quality=92)


def generate_assets() -> None:
    ASSET_DIR.mkdir(parents=True, exist_ok=True)
    variants = [
        ("Smart Login", "Secure access for every investor", "login", (37, 99, 235)),
        ("Market Dashboard", "Track the market in one glance", "dashboard", (14, 165, 233)),
        ("Stock Listing", "Discover gainers, losers, and sectors", "listing", (16, 185, 129)),
        ("Stock Details", "Understand charts before action", "details", (34, 197, 94)),
        ("Buy Order", "Fast and transparent order placement", "buy", (37, 99, 235)),
        ("Sell Order", "Exit positions with confidence", "sell", (220, 38, 38)),
        ("Add Funds", "Secure payment methods for trading", "funds_upi", (37, 99, 235)),
        ("Payment Method", "Choose the best transfer option", "funds_card", (124, 58, 237)),
        ("Funds Success", "Balance updated and ready", "funds_success", (22, 163, 74)),
        ("Order Confirmation", "Review every trade instantly", "confirmation", (37, 99, 235)),
        ("Portfolio View", "Monitor holdings and returns", "portfolio", (16, 185, 129)),
    ]
    for index, (title, subtitle, variant, accent) in enumerate(variants, start=2):
        create_screen_image(ASSET_DIR / f"image{index}.jpeg", title, subtitle, variant, accent)


def replace_paragraph_text(paragraph, new_text: str) -> None:
    if paragraph.runs:
        paragraph.runs[0].text = new_text
        for run in paragraph.runs[1:]:
            run.text = ""
    else:
        paragraph.add_run(new_text)


def set_cover_page_layout(doc: Document) -> None:
    for index in (14, 15, 16, 17):
        doc.paragraphs[index].alignment = WD_ALIGN_PARAGRAPH.CENTER

    section = doc.sections[0]
    usable_width = section.page_width.inches - section.left_margin.inches - section.right_margin.inches
    right_edge = Inches(usable_width)

    name_line = doc.paragraphs[18]
    name_line.alignment = WD_ALIGN_PARAGRAPH.LEFT
    replace_paragraph_text(name_line, "Narode sir\tProf. S.A.Patil")
    name_line.paragraph_format.tab_stops.add_tab_stop(right_edge, WD_TAB_ALIGNMENT.RIGHT)

    title_line = doc.paragraphs[19]
    title_line.alignment = WD_ALIGN_PARAGRAPH.LEFT
    replace_paragraph_text(title_line, "(Subject Teacher)\t(HOD CM)")
    title_line.paragraph_format.tab_stops.add_tab_stop(right_edge, WD_TAB_ALIGNMENT.RIGHT)


def update_document_text() -> None:
    doc = Document(str(SOURCE_DOC))
    for index, new_text in TEXT_REPLACEMENTS.items():
        replace_paragraph_text(doc.paragraphs[index], new_text)
    set_cover_page_layout(doc)
    doc.save(str(OUTPUT_DOC))


def get_source_logo() -> bytes:
    with zipfile.ZipFile(SOURCE_DOC, "r") as source:
        return source.read("word/media/image1.jpeg")


def replace_media() -> None:
    temp_doc = OUTPUT_DOC.with_name(f"{OUTPUT_DOC.stem}_temp.zip")
    source_zip = OUTPUT_DOC.with_suffix(".zip")

    if temp_doc.exists():
        temp_doc.unlink()
    if source_zip.exists():
        source_zip.unlink()

    shutil.copy2(OUTPUT_DOC, source_zip)
    with zipfile.ZipFile(source_zip, "r") as src, zipfile.ZipFile(temp_doc, "w", compression=zipfile.ZIP_DEFLATED) as dst:
        for item in src.infolist():
            name = item.filename
            if name.startswith("word/media/image") and name.endswith(".jpeg"):
                if name == "word/media/image1.jpeg":
                    dst.writestr(name, get_source_logo())
                else:
                    asset_path = ASSET_DIR / Path(name).name
                    dst.write(asset_path, arcname=name)
            else:
                dst.writestr(item, src.read(name))

    source_zip.unlink()
    if OUTPUT_DOC.exists():
        OUTPUT_DOC.unlink()
    temp_doc.rename(OUTPUT_DOC)


def verify_output() -> None:
    doc = Document(str(OUTPUT_DOC))
    checks = [
        doc.paragraphs[6].text,
        doc.paragraphs[11].text,
        doc.paragraphs[14].text,
        doc.paragraphs[17].text,
        doc.paragraphs[205].text,
    ]
    for item in checks:
        print(item)
    print(f"Inline shapes: {len(doc.inline_shapes)}")


def main() -> None:
    generate_assets()
    update_document_text()
    replace_media()
    verify_output()
    print(f"Created: {OUTPUT_DOC}")


if __name__ == "__main__":
    main()
