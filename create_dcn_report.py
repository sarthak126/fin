from __future__ import annotations

import zipfile
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont


SOURCE_DOC = Path(r"C:\Users\sarth\Downloads\microproject_UIUX (1).docx")
OUTPUT_DOC = Path(r"C:\Users\sarth\Downloads\Small_Office_Network_Design_DCN_Report.docx")
ASSET_DIR = Path(r"C:\Users\sarth\OneDrive\Desktop\ai\dcn_report_assets")
LOGO_PATH = ASSET_DIR / "institute_logo.jpeg"

FONT_REGULAR = Path(r"C:\Windows\Fonts\arial.ttf")
FONT_BOLD = Path(r"C:\Windows\Fonts\arialbd.ttf")

GROUP_MEMBERS = [
    "1. Sarthak Bhokare (26)",
    "2. Piyush Borse (29)",
    "3. Chine Roshan (33)",
    "4. Aditya Chavan (30)",
]


def font(size: int, bold: bool = False) -> ImageFont.FreeTypeFont | ImageFont.ImageFont:
    font_path = FONT_BOLD if bold else FONT_REGULAR
    try:
        return ImageFont.truetype(str(font_path), size=size)
    except OSError:
        return ImageFont.load_default()


def ensure_logo() -> None:
    ASSET_DIR.mkdir(parents=True, exist_ok=True)
    if LOGO_PATH.exists():
        return
    with zipfile.ZipFile(SOURCE_DOC, "r") as archive:
        LOGO_PATH.write_bytes(archive.read("word/media/image1.jpeg"))


def add_page_number(run) -> None:
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), "PAGE")
    run._r.append(fld)


def add_header_footer(section) -> None:
    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = header.add_run("Small Office Network Design - DCN")
    run.font.name = "Times New Roman"
    run.font.size = Pt(10)
    run.bold = True

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run("Page ")
    run.font.name = "Times New Roman"
    run.font.size = Pt(10)
    add_page_number(run)


def style_run(run, size: int = 12, bold: bool = False) -> None:
    run.font.name = "Times New Roman"
    run.font.size = Pt(size)
    run.bold = bold


def add_center_paragraph(doc: Document, text: str, size: int, bold: bool = False, space_after: int = 0):
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_after = Pt(space_after)
    run = paragraph.add_run(text)
    style_run(run, size=size, bold=bold)
    return paragraph


def add_heading(doc: Document, text: str):
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(4)
    paragraph.paragraph_format.space_after = Pt(6)
    run = paragraph.add_run(text)
    style_run(run, size=14, bold=True)
    run.font.color.rgb = RGBColor(24, 47, 83)
    return paragraph


def add_body(doc: Document, text: str, spacing: int = 4):
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    paragraph.paragraph_format.line_spacing = 1.15
    paragraph.paragraph_format.space_after = Pt(spacing)
    run = paragraph.add_run(text)
    style_run(run, size=12)
    return paragraph


def add_line(doc: Document, text: str):
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.left_indent = Inches(0.2)
    paragraph.paragraph_format.space_after = Pt(2)
    run = paragraph.add_run(text)
    style_run(run, size=12)
    return paragraph


def add_caption(doc: Document, text: str):
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(2)
    paragraph.paragraph_format.space_after = Pt(6)
    run = paragraph.add_run(text)
    style_run(run, size=10, bold=True)
    return paragraph


def set_cell_text(cell, text: str, bold: bool = False, center: bool = False, size: int = 11):
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if center else WD_ALIGN_PARAGRAPH.LEFT
    run = paragraph.add_run(text)
    style_run(run, size=size, bold=bold)


def create_topology_image(path: Path) -> None:
    width, height = 1500, 900
    image = Image.new("RGB", (width, height), (246, 250, 255))
    draw = ImageDraw.Draw(image)

    draw.rounded_rectangle((0, 0, width, 120), radius=0, fill=(20, 54, 94))
    draw.text((60, 34), "Small Office Network Topology", font=font(42, bold=True), fill="white")
    draw.text((60, 82), "DCN project layout for 10-20 computers using router and switch", font=font(20), fill=(221, 235, 249))

    def box(x1, y1, x2, y2, label, fill, outline=(18, 42, 72)):
        draw.rounded_rectangle((x1, y1, x2, y2), radius=28, fill=fill, outline=outline, width=4)
        bbox = draw.multiline_textbbox((0, 0), label, font=font(24, bold=True), spacing=6)
        tx = x1 + ((x2 - x1) - (bbox[2] - bbox[0])) / 2
        ty = y1 + ((y2 - y1) - (bbox[3] - bbox[1])) / 2
        draw.multiline_text((tx, ty), label, font=font(24, bold=True), fill=(17, 24, 39), align="center", spacing=6)

    def connect(p1, p2, color=(37, 99, 235), width_px=8):
        draw.line((p1, p2), fill=color, width=width_px)

    box(610, 170, 900, 280, "Internet", (220, 236, 255))
    box(610, 340, 900, 460, "Router", (209, 250, 229))
    box(565, 530, 945, 665, "24-Port Switch", (254, 243, 199))

    connect((755, 280), (755, 340))
    connect((755, 460), (755, 530))

    department_boxes = [
        (80, 720, 320, 840, "Accounts\n4 PCs"),
        (360, 720, 600, 840, "Admin\n3 PCs"),
        (640, 720, 880, 840, "Support\n4 PCs"),
        (920, 720, 1160, 840, "Sales\n3 PCs"),
        (1200, 720, 1420, 840, "Printer +\nWi-Fi AP"),
    ]
    colors = [(239, 246, 255), (243, 232, 255), (236, 253, 245), (255, 237, 213), (254, 242, 242)]

    for idx, (x1, y1, x2, y2, label) in enumerate(department_boxes):
        box(x1, y1, x2, y2, label, colors[idx])
        connect((755, 665), ((x1 + x2) // 2, y1), width_px=6)

    draw.rounded_rectangle((84, 170, 430, 450), radius=28, fill=(255, 255, 255), outline=(191, 219, 254), width=4)
    draw.text((114, 210), "Design Notes", font=font(28, bold=True), fill=(30, 64, 175))
    notes = [
        "15 office computers",
        "Single router for internet access",
        "One managed switch for LAN",
        "Shared printer and Wi-Fi access point",
        "Star topology for easy maintenance",
    ]
    for idx, line in enumerate(notes):
        draw.text((118, 268 + idx * 34), f"- {line}", font=font(20), fill=(31, 41, 55))

    image.save(path, quality=92)


def create_layout_image(path: Path) -> None:
    width, height = 1500, 880
    image = Image.new("RGB", (width, height), (251, 252, 255))
    draw = ImageDraw.Draw(image)

    draw.rounded_rectangle((40, 40, 1460, 840), radius=36, outline=(31, 41, 55), width=6, fill=(255, 255, 255))
    draw.text((70, 72), "Proposed Small Office Layout", font=font(40, bold=True), fill=(24, 47, 83))
    draw.text((72, 122), "Simple floor arrangement showing departments and network points", font=font(21), fill=(75, 85, 99))

    rooms = [
        (90, 180, 450, 420, "Accounts Section\n4 Computers", (239, 246, 255)),
        (500, 180, 860, 420, "Admin Cabin\n3 Computers", (237, 233, 254)),
        (910, 180, 1410, 420, "Sales & Support\n7 Computers", (236, 253, 245)),
        (90, 470, 520, 760, "Reception + Waiting", (255, 247, 237)),
        (570, 470, 940, 760, "Network Rack Area\nRouter + Switch + UPS", (254, 242, 242)),
        (990, 470, 1410, 760, "Printer / Meeting Area\nPrinter + Wi-Fi AP", (254, 249, 195)),
    ]

    for x1, y1, x2, y2, label, fill in rooms:
        draw.rounded_rectangle((x1, y1, x2, y2), radius=24, fill=fill, outline=(148, 163, 184), width=3)
        bbox = draw.multiline_textbbox((0, 0), label, font=font(28, bold=True), spacing=8)
        tx = x1 + ((x2 - x1) - (bbox[2] - bbox[0])) / 2
        ty = y1 + ((y2 - y1) - (bbox[3] - bbox[1])) / 2
        draw.multiline_text((tx, ty), label, font=font(28, bold=True), fill=(17, 24, 39), align="center", spacing=8)

    draw.line((755, 760, 755, 840), fill=(37, 99, 235), width=6)
    draw.text((620, 794), "Structured Cabling Route", font=font(22, bold=True), fill=(37, 99, 235))
    image.save(path, quality=92)


def build_cover_page(doc: Document) -> None:
    add_center_paragraph(doc, "Sanjivani Rural Education Society's", 15, bold=True, space_after=2)
    add_center_paragraph(doc, "Sanjivani K.B.P Polytechnic", 15, bold=True, space_after=2)
    add_center_paragraph(doc, "DEPARTMENT OF COMPUTER TECHNOLOGY", 14, bold=True, space_after=2)
    add_center_paragraph(doc, "(2025-26)", 12, bold=True, space_after=8)

    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    run.add_picture(str(LOGO_PATH), width=Inches(1.65))
    paragraph.paragraph_format.space_after = Pt(8)

    add_center_paragraph(doc, "A PROJECT REPORT ON", 13, bold=True, space_after=6)
    add_center_paragraph(doc, '"Small Office Network Design"', 16, bold=True, space_after=2)
    add_center_paragraph(doc, "(10-20 Computers with Router and Switch)", 12, bold=True, space_after=8)
    add_center_paragraph(doc, "SUBMITTED TO THE MSBTE, MUMBAI", 12, bold=True, space_after=2)
    add_center_paragraph(doc, "Under 'DCN'", 12, bold=True, space_after=10)
    add_center_paragraph(doc, "Under The Guidance", 12, bold=True, space_after=2)
    add_center_paragraph(doc, "Sangle sir", 12, bold=True, space_after=10)
    add_center_paragraph(doc, "SUBMITTED BY :", 12, bold=True, space_after=2)

    for member in GROUP_MEMBERS:
        add_center_paragraph(doc, member, 12, bold=False, space_after=1)

    name_line = doc.add_paragraph()
    name_line.alignment = WD_ALIGN_PARAGRAPH.LEFT
    name_line.paragraph_format.tab_stops.add_tab_stop(Inches(6.25), WD_TAB_ALIGNMENT.RIGHT)
    run = name_line.add_run("Sangle sir\tProf. S.A.Patil")
    style_run(run, size=12, bold=True)
    name_line.paragraph_format.space_before = Pt(10)

    designation = doc.add_paragraph()
    designation.alignment = WD_ALIGN_PARAGRAPH.LEFT
    designation.paragraph_format.tab_stops.add_tab_stop(Inches(6.25), WD_TAB_ALIGNMENT.RIGHT)
    run = designation.add_run("(Subject Teacher)\t(HOD CM)")
    style_run(run, size=11, bold=False)

    doc.add_page_break()


def build_content(doc: Document, topology_path: Path, layout_path: Path) -> None:
    add_heading(doc, "INTRODUCTION")
    add_body(
        doc,
        "The Small Office Network Design project focuses on creating a reliable and easy-to-manage local "
        "area network for a small office with 10 to 20 computers. The proposed design uses one router for "
        "internet connectivity and security control, one switch for internal communication, and structured "
        "cabling to connect all end devices. The aim of the project is to provide stable resource sharing, "
        "simple maintenance, and secure day-to-day office communication.",
    )

    add_heading(doc, "OBJECTIVES OF THE PROJECT")
    objectives = [
        "To design a simple and scalable office network for 10 to 20 computers.",
        "To provide internet access, printer sharing, and easy communication among departments.",
        "To use a router and switch based star topology for efficient management.",
        "To reduce downtime and simplify troubleshooting in the office environment.",
        "To create a design that can be expanded in the future without major changes.",
    ]
    for item in objectives:
        add_line(doc, f"- {item}")

    add_heading(doc, "COMPONENTS USED")
    components = doc.add_table(rows=1, cols=3)
    components.alignment = WD_TABLE_ALIGNMENT.CENTER
    components.style = "Table Grid"
    headers = ["Component", "Quantity", "Purpose"]
    for idx, text in enumerate(headers):
        set_cell_text(components.rows[0].cells[idx], text, bold=True, center=True, size=11)

    rows = [
        ("Router", "1", "Connects office LAN to the internet and controls traffic"),
        ("24-Port Switch", "1", "Provides LAN connectivity to all computers"),
        ("Desktop Computers", "15", "Used by office staff in different departments"),
        ("Printer", "1", "Shared printing resource for the office"),
        ("Wi-Fi Access Point", "1", "Supports wireless access for mobile devices"),
        ("Cat6 Cables", "As required", "Provides reliable structured cabling"),
    ]
    for component, quantity, purpose in rows:
        cells = components.add_row().cells
        set_cell_text(cells[0], component, size=10)
        set_cell_text(cells[1], quantity, center=True, size=10)
        set_cell_text(cells[2], purpose, size=10)

    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    run.add_picture(str(layout_path), width=Inches(5.8))
    add_caption(doc, "Fig. 1 Proposed office layout and network equipment placement")

    doc.add_page_break()

    add_heading(doc, "NETWORK TOPOLOGY")
    add_body(
        doc,
        "The network follows a star topology in which all computers, the printer, and the wireless access "
        "point are connected to a central switch. The switch is then connected to the router, which manages "
        "internet access for the entire office. This topology is preferred because it is easy to maintain, "
        "supports fault isolation, and allows future expansion when more systems are added.",
    )

    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    run.add_picture(str(topology_path), width=Inches(6.2))
    add_caption(doc, "Fig. 2 Small office star topology using router and switch")

    add_body(
        doc,
        "In this design, departments such as accounts, admin, support, and sales remain logically connected "
        "through the same switch. Shared devices like the network printer can be accessed by all authorized "
        "users, and the router can be configured with DHCP, firewall rules, and basic access control.",
    )

    doc.add_page_break()

    add_heading(doc, "IP ADDRESSING PLAN")
    ip_table = doc.add_table(rows=1, cols=4)
    ip_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    ip_table.style = "Table Grid"
    headers = ["Device / Area", "Count", "IP Range", "Remarks"]
    for idx, text in enumerate(headers):
        set_cell_text(ip_table.rows[0].cells[idx], text, bold=True, center=True, size=11)

    ip_rows = [
        ("Router LAN", "1", "192.168.1.1", "Default gateway"),
        ("Accounts", "4", "192.168.1.10 - 13", "Static or DHCP reservation"),
        ("Admin", "3", "192.168.1.20 - 22", "Department systems"),
        ("Support", "4", "192.168.1.30 - 33", "Shared support team"),
        ("Sales", "4", "192.168.1.40 - 43", "Sales staff computers"),
        ("Printer", "1", "192.168.1.50", "Static IP"),
        ("Wi-Fi Devices", "Dynamic", "192.168.1.100 - 120", "DHCP pool"),
    ]
    for area, count, ip_range, remarks in ip_rows:
        cells = ip_table.add_row().cells
        set_cell_text(cells[0], area, size=10)
        set_cell_text(cells[1], count, center=True, size=10)
        set_cell_text(cells[2], ip_range, center=True, size=10)
        set_cell_text(cells[3], remarks, size=10)

    add_heading(doc, "WORKING OF THE SYSTEM")
    steps = [
        "The internet connection enters the office through the router.",
        "The router provides IP addresses and forwards local traffic to the switch.",
        "The switch connects all computers, the printer, and the Wi-Fi access point.",
        "Users can share files, access the printer, and browse the internet through the same network.",
        "If one computer fails, the rest of the network continues to operate normally.",
    ]
    for idx, item in enumerate(steps, start=1):
        add_line(doc, f"{idx}. {item}")

    add_heading(doc, "SECURITY FEATURES")
    for item in [
        "Router firewall rules to block unwanted external traffic",
        "Strong Wi-Fi password and restricted admin access",
        "Antivirus and OS updates on all computers",
        "Regular cable checks and backup of important office data",
    ]:
        add_line(doc, f"- {item}")

    doc.add_page_break()

    add_heading(doc, "ADVANTAGES OF THE SYSTEM")
    for item in [
        "Simple and cost-effective network for a small office",
        "Easy resource sharing between departments",
        "Centralized connection through router and switch",
        "Quick troubleshooting due to star topology",
        "Future expansion is possible without redesigning the whole network",
    ]:
        add_line(doc, f"- {item}")

    add_heading(doc, "LIMITATIONS OF THE SYSTEM")
    for item in [
        "The design depends on the switch as a central device",
        "No advanced enterprise redundancy is included",
        "Internet performance depends on the router and service provider",
        "Security level is basic compared to larger corporate networks",
    ]:
        add_line(doc, f"- {item}")

    add_heading(doc, "FUTURE ENHANCEMENTS")
    for item in [
        "VLAN-based department separation",
        "Centralized server and file backup system",
        "Biometric access for server room or rack area",
        "Improved monitoring using network management tools",
    ]:
        add_line(doc, f"- {item}")

    add_heading(doc, "CONCLUSION")
    add_body(
        doc,
        "The Small Office Network Design project presents a practical DCN solution for a growing office "
        "environment. By using one router, one switch, and a structured star topology, the design offers "
        "easy management, reliable connectivity, and sufficient scalability for 10 to 20 computers. This "
        "network model is suitable for day-to-day office work and can be enhanced further as the organization expands.",
    )


def main() -> None:
    ensure_logo()
    topology_path = ASSET_DIR / "topology.jpeg"
    layout_path = ASSET_DIR / "office_layout.jpeg"
    create_topology_image(topology_path)
    create_layout_image(layout_path)

    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    add_header_footer(section)

    build_cover_page(doc)

    content_section = doc.sections[-1]
    content_section.start_type = WD_SECTION.NEW_PAGE
    add_header_footer(content_section)

    build_content(doc, topology_path, layout_path)
    doc.save(str(OUTPUT_DOC))

    verify = Document(str(OUTPUT_DOC))
    print(f"Created: {OUTPUT_DOC}")
    print(f"Paragraphs: {len(verify.paragraphs)}")
    print(f"Tables: {len(verify.tables)}")
    print(f"Images: {len(verify.inline_shapes)}")


if __name__ == "__main__":
    main()
